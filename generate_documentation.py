#!/usr/bin/env python3
"""
Generate Design Documentation for Spectra Media Agent
Outputs in PDF, DOCX, and RTF formats
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from datetime import datetime

# ============================================================================
# DESIGN DOCUMENTATION CONTENT
# ============================================================================

TITLE = "Spectra Media Agent - Design Documentation"
DATE_CREATED = datetime.now().strftime("%B %d, %Y")
VERSION = "1.0"

# Define all sections
SECTIONS = {
    "executive_summary": {
        "title": "1. EXECUTIVE SUMMARY",
        "content": """
Spectra Media Agent is an advanced agentic platform that automates the creation, deployment, and optimization of marketing campaigns across multiple advertising platforms including Google Ads and Facebook Ads. The system leverages artificial intelligence agents to autonomously manage complex marketing workflows while maintaining human oversight and control.

Key Capabilities:
• Autonomous strategy generation based on business inputs
• Automated campaign creation and deployment across multiple platforms
• Real-time performance monitoring and analysis
• AI-driven optimization and recommendations
• Multi-platform budget allocation and management
• Integration with Stripe billing for cost tracking
• Role-based access control and permission management
"""
    },
    "system_architecture": {
        "title": "2. SYSTEM ARCHITECTURE",
        "content": """
2.1 HIGH-LEVEL ARCHITECTURE

Spectra Media Agent follows a distributed agent-based architecture:

┌─────────────────────────────────────────────────────────────┐
│                      User Interface (Vue.js)                 │
└────────────────────────┬────────────────────────────────────┘
                         │
┌────────────────────────▼────────────────────────────────────┐
│               API Layer (Laravel/HTTP)                      │
├─────────────────────────────────────────────────────────────┤
│  Authentication │ Campaigns │ Performance │ Admin │ Billing  │
└────────────────────────┬────────────────────────────────────┘
                         │
┌────────────────────────▼────────────────────────────────────┐
│           Core Services & Business Logic                     │
├─────────────────────────────────────────────────────────────┤
│ Strategy   │ Performance  │ Campaign    │ Billing │ Reports  │
│ Generator  │ Analyzer     │ Manager     │ Service │ Service   │
└────────────────────────┬────────────────────────────────────┘
                         │
┌────────────────────────▼────────────────────────────────────┐
│          Agent Jobs (Queue-based Execution)                  │
├─────────────────────────────────────────────────────────────┤
│ Strategy   │ Fetch       │ Deploy      │ Optimize │ Reporting│
│ Generation │ Performance │ Campaigns   │ Campaigns│ Agents   │
└────────────────────────┬────────────────────────────────────┘
                         │
┌────────────────────────▼────────────────────────────────────┐
│         External Platform Integrations                       │
├─────────────────────────────────────────────────────────────┤
│ Google Ads API │ Facebook Ads API │ Stripe API │ Gemini API  │
└─────────────────────────────────────────────────────────────┘

2.2 CORE COMPONENTS

DATABASE LAYER:
- PostgreSQL with vector extensions (pgvector)
- Stores campaigns, performance data, strategies, and audit trails
- Supports multi-tenant architecture with customer isolation

QUEUE SYSTEM:
- Redis-based job queue (via Horizon)
- Enables asynchronous agent execution
- Supports job retries, timeouts, and monitoring

CACHE LAYER:
- Redis cache for performance optimization
- Caches API responses and computed data
"""
    },
    "agent_design": {
        "title": "3. AGENT DESIGN & WORKFLOW",
        "content": """
3.1 PRIMARY AGENTS

STRATEGY GENERATION AGENT
Purpose: Generate and refine marketing strategies
Implementation: GenerateStrategy Job + StrategyGenerator Service
Input: Business details, performance data, seasonal adjustments
Output: Platform-specific strategies with bidding parameters
AI Integration: Uses Google Gemini 3.x API for strategy synthesis

PERFORMANCE ANALYSIS AGENT
Purpose: Collect, analyze, and report on campaign metrics
Implementation: FetchGoogleAdsPerformanceData, FetchFacebookAdsPerformanceData
Input: Campaign IDs, date ranges
Output: Aggregated performance metrics with insights
Features: Raw metric collection, currency conversion, anomaly detection

CAMPAIGN DEPLOYMENT AGENT
Purpose: Create and launch campaigns on advertising platforms
Implementation: DeployCampaignToGoogleAds, DeployCampaignToFacebook
Input: Campaign specifications, creative assets
Output: Campaign creation confirmations with platform IDs
Features: Automated asset upload, validation, platform-specific optimization

OPTIMIZATION AGENT
Purpose: Recommend and apply optimizations based on performance
Implementation: GenerateOptimizationRecommendations, ApplyOptimizations
Input: Performance metrics, historical data
Output: Optimization recommendations with implementation details
Features: A/B testing suggestions, budget reallocation, bidding adjustments

3.2 AGENT WORKFLOW

User Input
    ↓
Strategy Generation → Strategy Approval
    ↓
Campaign Creation → Asset Upload
    ↓
Campaign Deployment → Platform Sync
    ↓
Performance Monitoring ← Continuous Fetching
    ↓
Analysis & Recommendations
    ↓
Optimization Implementation
    ↓
Feedback Loop (repeat)

3.3 JOB QUEUE INTEGRATION

All agents are implemented as Laravel Jobs dispatched through the queue system:
- Asynchronous execution prevents blocking user requests
- Retry logic with exponential backoff for reliability
- Circuit breaker pattern for external API calls
- Distributed locking to prevent concurrent conflict
- Comprehensive error logging and alerting
"""
    },
    "data_flow": {
        "title": "4. DATA FLOW & INTEGRATION",
        "content": """
4.1 CAMPAIGN CREATION FLOW

1. User Input (via Web UI)
   - Business objective
   - Target audience
   - Budget parameters
   - Timeline

2. Strategy Generation
   - AI analyzes business inputs
   - Generates platform-specific strategies
   - Creates ad copy variations
   - Defines imagery/video recommendations
   - Specifies bidding strategy

3. User Approval
   - Review generated strategies
   - Approve or request modifications
   - Define final budget allocation

4. Campaign Creation
   - Create campaign structure
   - Upload creative assets
   - Configure targeting parameters
   - Set budget and scheduling

5. Platform Deployment
   - Deploy to Google Ads
   - Deploy to Facebook Ads
   - Sync campaign IDs back to database

6. Monitoring & Optimization
   - Fetch performance metrics (every 2 hours)
   - Analyze performance trends
   - Generate optimization recommendations
   - Apply approved optimizations

4.2 DATA MODELS

KEY TABLES:
- customers: User accounts and platform credentials
- campaigns: Campaign definitions and metadata
- campaign_assets: Creative assets (images, videos, copy)
- google_ads_campaigns: Google Ads specific data
- facebook_ads_campaigns: Facebook Ads specific data
- strategies: AI-generated strategies with approval status
- performance_data: Aggregated metrics from platforms
- billing_records: Cost tracking and invoicing data
- optimization_recommendations: AI suggestions for improvements

4.3 API INTEGRATIONS

GOOGLE ADS API:
- Authentication: OAuth 2.0 with refresh token
- Capabilities: Campaign CRUD, performance reporting, asset management
- Rate Limiting: Handled by Google
- Error Handling: Exponential backoff, retry logic

FACEBOOK GRAPH API:
- Authentication: App tokens with user delegation
- Capabilities: Campaign CRUD, insights/metrics, creative upload
- Rate Limiting: Throttling and request batching
- Error Handling: Automatic retry with exponential backoff

GOOGLE GEMINI API:
- Purpose: Strategic decision making and analysis
- Models: Gemini 3.x (Flash for speed, Pro for quality)
- Features: Vision capabilities for image analysis, text generation
- Safety: Content filtering enabled by default
"""
    },
    "security": {
        "title": "5. SECURITY & COMPLIANCE",
        "content": """
5.1 AUTHENTICATION & AUTHORIZATION

Multi-layered authentication approach:
- Laravel Sanctum for API authentication
- JWT tokens for session management
- OAuth 2.0 for third-party platform integration
- Role-based access control (RBAC)
  - Admin: Full system access
  - Manager: Campaign and strategy management
  - Viewer: Read-only access
  - Analyst: Performance analysis and reporting

5.2 DATA PROTECTION

Encryption Standards:
- All platform credentials encrypted at rest using AES-256
- In-transit TLS 1.2+ for all network communications
- Database encryption for sensitive data fields
- No plaintext storage of API keys or tokens

credential Management:
- Refresh tokens encrypted before storage
- Automatic token rotation on use
- Secure deletion of revoked credentials
- Audit logging of credential access

5.3 AUDIT & COMPLIANCE

Complete audit trail:
- All user actions logged with timestamps
- Campaign modifications tracked with before/after states
- Performance metric changes recorded
- API call logging for debugging
- GDPR compliance features:
  - Data export functionality
  - Right to deletion with cascading deletes
  - Consent tracking for third-party integrations

5.4 PLATFORM-SPECIFIC REQUIREMENTS

Google Ads:
- Developer token management
- MCC account support for enterprise
- Scopes: adwords (campaign management), tagmanager.edit.containers

Facebook Ads:
- App review process compliance
- Business verification requirements
- Ad library transparency
- Conversion tracking compliance
"""
    },
    "deployment": {
        "title": "6. DEPLOYMENT & INFRASTRUCTURE",
        "content": """
6.1 DEPLOYMENT ARCHITECTURE

Production Deployment:
- Hosted on Laravel Forge
- Automated deployment via Git
- Load balanced for high availability
- SSL certificates via Let's Encrypt
- Auto-scaling capabilities

Technology Stack:
- Web Server: Nginx
- Application: Laravel 11.x (PHP 8.3+)
- Database: PostgreSQL with pgvector
- Cache: Redis
- Queue: Horizon (Redis-backed)
- Frontend: Vue.js 3.x with Inertia.js
- Build Tool: Vite for frontend assets

6.2 ENVIRONMENT CONFIGURATION

Production Environment Variables:
- API credentials for Google, Facebook, Stripe
- Database connection details
- Queue and cache configuration
- Mail service configuration
- Logging and monitoring setup
- Feature flags for beta features

6.3 SCALING CONSIDERATIONS

Load Balancing:
- Horizontal scaling via Forge
- Database connection pooling
- Redis cluster for caching/queues
- CDN for frontend assets

Queue Processing:
- Multiple worker processes for job processing
- Timeout configuration for long-running jobs
- Dead-letter queue for failed jobs
- Circuit breaker for external APIs

Monitoring & Alerting:
- Error tracking with Sentry
- Performance monitoring
- Log aggregation
- Custom health checks
"""
    },
    "maintenance": {
        "title": "7. MAINTENANCE & OPERATIONS",
        "content": """
7.1 REGULAR MAINTENANCE TASKS

Daily:
- Monitor queue health and failed job counts
- Check error logs for critical issues
- Verify API integrations operational
- Monitor database performance

Weekly:
- Review performance trends
- Analyze campaign optimization recommendations
- Audit user access and permissions
- Check backup integrity

Monthly:
- Database optimization and maintenance
- Dependency updates (with testing)
- Security patches
- Cost analysis and billing reconciliation

7.2 MONITORING & ALERTING

Key Metrics:
- Queue depth and processing time
- API response times
- Database query performance
- Cache hit rates
- Error rates by component
- Campaign performance variance

Alert Thresholds:
- Queue depth > 10,000 jobs
- API response time > 5 seconds
- Database CPU > 80%
- Campaign underperformance triggers
- Failed job count increasing

7.3 BACKUP & DISASTER RECOVERY

Backup Strategy:
- Daily automated database backups
- 30-day retention policy
- Off-site backup replication
- Point-in-time recovery capability
- Regular restore testing

Disaster Recovery:
- RTO (Recovery Time Objective): 1 hour
- RPO (Recovery Point Objective): 1 hour
- Documented runbooks for common issues
- Instance redundancy in production
- Database failover configuration

7.4 DEPENDENCY MANAGEMENT

Regular Updates:
- Laravel framework updates (quarterly)
- PHP minor versions (quarterly)
- Security-critical patches (immediate)
- Development dependencies
- Database extensions

Testing Protocol:
- Automated test suite execution
- Staging environment validation
- Smoke tests post-deployment
- Performance regression testing
"""
    },
    "performance": {
        "title": "8. PERFORMANCE & OPTIMIZATION",
        "content": """
8.1 PERFORMANCE TARGETS

API Response Times:
- Campaign list: < 500ms
- Strategy generation: < 5000ms (async)
- Performance reports: < 2000ms
- Campaign deployment: < 3000ms

Database Performance:
- Query execution: < 100ms (95th percentile)
- Connection pool: 20-50 active connections
- Slow query threshold: > 500ms

Frontend Performance:
- Page load: < 3 seconds
- Interactive elements: < 1 second response
- Graph rendering: < 500ms

8.2 OPTIMIZATION STRATEGIES

Query Optimization:
- Selective column selection via projections
- Composite indexing on common filters
- Query result caching with TTL
- N+1 query elimination through eager loading

Caching Architecture:
- API response caching (5 min TTL)
- Computed metrics caching (hourly TTL)
- User preference caching
- Strategy caching until approval

Job Batch Processing:
- Campaign creation in batches
- Performance metric fetching in batches
- Bulk operations for cost efficiency
- Progressive status updates

8.3 SCALABILITY DESIGN

Horizontal Scaling:
- Stateless application design
- Load balancer distribution
- Message queue for async jobs
- Distributed caching

Vertical Scaling:
- Database server resources (CPU, RAM)
- Cache server sizing
- Connection pool optimization
- Job worker allocation

8.4 LOAD TESTING RESULTS

Expected Capacity:
- 10,000+ concurrent users
- 100+ campaigns per account
- Real-time data sync for 1,000+ campaigns
- 1,000+ jobs/minute queue throughput
"""
    },
    "future_development": {
        "title": "9. FUTURE DEVELOPMENT & ROADMAP",
        "content": """
9.1 PLANNED FEATURES

Q2 2026:
- Microsoft Ads integration
- Advanced A/B testing framework
- Predictive analytics for campaign performance
- Automated bid adjustment policies

Q3 2026:
- TikTok Ads integration
- LinkedIn Ads integration
- Advanced audience segmentation
- Predictive budget allocation

Q4 2026:
- Marketplace expansion capabilities
- Custom model training pipeline
- Advanced attribution modeling
- API for partner integrations

9.2 TECHNICAL IMPROVEMENTS

Infrastructure:
- Migration to Kubernetes for container orchestration
- Multi-region deployment for global coverage
- Enhanced disaster recovery with active-active setup
- Database sharding for ultra-scale

Architecture:
- Event-driven architecture with event sourcing
- GraphQL API layer (beta)
- Microservices decomposition (selective)
- Enhanced real-time capabilities with WebSockets

AI/ML Enhancements:
- Fine-tuned models for specific industries
- Real-time sentiment analysis
- Competitive intelligence integration
- Automated creative generation

9.3 COMMUNITY & ECOSYSTEM

- Developer API for third-party integrations
- Public documentation and SDKs
- Community plugin marketplace
- Open-source component libraries

9.4 RESEARCH & INNOVATION

Exploring:
- Advanced multi-armed bandit optimization
- Causal inference for attribution
- Federated learning for privacy
- Reinforcement learning for bidding strategies
"""
    },
    "appendix": {
        "title": "10. APPENDIX - TECHNICAL REFERENCE",
        "content": """
10.1 PROJECT STRUCTURE

app/
├── Console/              # CLI commands
├── Http/                 # Controllers and requests
├── Jobs/                 # Queue-based agents
├── Models/               # Eloquent models
├── Services/             # Business logic
│   ├── GoogleAds/       # Google Ads integration
│   ├── FacebookAds/      # Facebook Ads integration
│   ├── Billing/         # Payment processing
│   └── AI/              # LLM integrations
├── Observers/           # Model event listeners
├── Policies/            # Authorization policies
└── Exceptions/          # Custom exceptions

config/
├── googleads.php        # Google Ads configuration
├── campaigns.php        # Campaign settings
├── budget_rules.php     # Budget allocation rules
├── seasonal_strategies.php
└── platform_rules.php   # Platform-specific rules

database/
├── migrations/          # Database schema
├── factories/          # Model factories for testing
└── seeders/            # Database seeders

resources/
├── js/                  # Vue.js frontend
├── views/               # Blade templates
└── css/                 # Styling

10.2 KEY FILES & RESPONSIBILITIES

Strategy Management:
- app/Services/StrategyGenerator.php
- app/Jobs/GenerateStrategy.php
- app/Models/Strategy.php

Campaign Management:
- app/Services/CampaignManager.php
- app/Models/Campaign.php
- app/Http/Controllers/CampaignController.php

Performance Tracking:
- app/Services/PerformanceAnalyzer.php
- app/Jobs/FetchGoogleAdsPerformanceData.php
- app/Jobs/FetchFacebookAdsPerformanceData.php

Billing:
- app/Services/Billing/BillingService.php
- app/Models/BillingRecord.php

10.3 EXTENSION POINTS

Custom Strategy Generators:
- Implement StrategyGeneratorInterface
- Register in service container
- Configure via platform_rules.php

Custom Platform Support:
- Create new PlatformService class
- Implement CampaignDeployerInterface
- Add performance data fetcher

Custom Optimization Rules:
- Implement OptimizationRuleInterface
- Register in optimization engine
- Configure thresholds and parameters

10.4 TESTING STRATEGY

Unit Tests:
- Service layer logic
- Model mutations and relations
- Utility functions

Feature Tests:
- API endpoints
- Workflow integration
- Database transactions

Integration Tests:
- API mocking for external services
- End-to-end workflow validation
- Database state verification

Performance Tests:
- Load testing with Apache JMeter
- Stress testing job queue
- Database query optimization validation
"""
    }
}

# ============================================================================
# DOCUMENT GENERATORS
# ============================================================================

def create_pdf(output_path):
    """Generate PDF documentation using ReportLab"""
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 51, 102),
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=RGBColor(0, 51, 102),
        spaceAfter=12,
        spaceBefore=12
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    # Title page
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph(TITLE, title_style))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(f"Version {VERSION}", styles['Normal']))
    story.append(Paragraph(f"Created: {DATE_CREATED}", styles['Normal']))
    story.append(PageBreak())
    
    # Table of Contents
    story.append(Paragraph("TABLE OF CONTENTS", heading_style))
    story.append(Spacer(1, 0.2*inch))
    for key, section in SECTIONS.items():
        story.append(Paragraph(section['title'], styles['Normal']))
    story.append(PageBreak())
    
    # Content sections
    for key, section in SECTIONS.items():
        story.append(Paragraph(section['title'], heading_style))
        for line in section['content'].strip().split('\n'):
            if line.strip():
                story.append(Paragraph(line, body_style))
        story.append(PageBreak())
    
    doc.build(story)
    print(f"✓ PDF generated: {output_path}")

def create_docx(output_path):
    """Generate DOCX documentation using python-docx"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(TITLE, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Version: {VERSION}")
    doc.add_paragraph(f"Created: {DATE_CREATED}")
    doc.add_page_break()
    
    # Add table of contents
    doc.add_heading("TABLE OF CONTENTS", level=1)
    for key, section in SECTIONS.items():
        doc.add_paragraph(section['title'], style='List Number')
    doc.add_page_break()
    
    # Add sections
    for key, section in SECTIONS.items():
        doc.add_heading(section['title'], level=1)
        content = section['content'].strip()
        for paragraph_text in content.split('\n'):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text)
                # Format special lines
                if any(x in paragraph_text for x in ['├─', '│', '└─', '─', '┌', '┐', '┘', '┴', '┬']):
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
        doc.add_page_break()
    
    doc.save(output_path)
    print(f"✓ Word document generated: {output_path}")

def create_rtf(output_path):
    """Generate RTF documentation"""
    rtf_content = r"{\rtf1\ansi\ansicpg1252\cocoartf1" + "\n"
    rtf_content += r"{\colortbl;\red255\green255\blue255;}" + "\n"
    rtf_content += r"{\*\expandedcolortbl;;}" + "\n"
    rtf_content += r"{\fonttbl\f0\fswiss Helvetica;}" + "\n"
    rtf_content += r"{\colortbl;\red255\green255\blue255;}" + "\n"
    rtf_content += r"{\*\expandedcolortbl;;}" + "\n"
    rtf_content += r"\margl1440\margr1440\margtsxn0\margbsxn0\vieww10800\viewh8400\viewkind0" + "\n"
    rtf_content += r"\pard\tx720\tx1440\tx2160\tx2880\tx3600\tx4320\tx5040\tx5760\tx6480\tx7200\tx7920\tx8640\pardirnatural\partightenfactor100" + "\n"
    rtf_content += "\n"
    
    # Title
    rtf_content += r"{\*\fldinst HYPERLINK \l "+'"'+"bookmark1"+'"'+"}{\\fldrslt " + TITLE + "}" + "\n"
    rtf_content += r"{\b\fs48 " + TITLE + r"}" + "\n\par" + "\n"
    rtf_content += r"{\fs24 Version " + VERSION + r"}" + "\n\par" + "\n"
    rtf_content += r"{\fs24 Created: " + DATE_CREATED + r"}" + "\n\par" + "\n\n"
    
    # Table of Contents
    rtf_content += r"{\b\fs32 TABLE OF CONTENTS}" + "\n\par\n"
    for key, section in SECTIONS.items():
        rtf_title = section['title'].replace('&', r"\\'26")
        rtf_content += r"{\fs20 " + rtf_title + r"}" + "\n\par\n"
    
    # Content sections
    for key, section in SECTIONS.items():
        section_title = section['title'].replace('&', r"\\'26")
        rtf_content += r"\page" + "\n"
        rtf_content += r"{\b\fs32 " + section_title + r"}" + "\n\par\n"
        
        for line in section['content'].strip().split('\n'):
            if line.strip():
                clean_line = line.replace('&', r"\\'26").replace('{', r"\{").replace('}', r"\}")
                rtf_content += r"{\fs20 " + clean_line + r"}" + "\n\par\n"
    
    rtf_content += "}"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(rtf_content)
    
    print(f"✓ RTF document generated: {output_path}")

# ============================================================================
# MAIN EXECUTION
# ============================================================================

if __name__ == "__main__":
    import os
    
    # Output directory
    output_dir = "/Users/josh/Documents/Github/spectra-media-agent/docs"
    
    # Ensure directory exists
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate documents
    pdf_path = os.path.join(output_dir, "Spectra_Media_Agent_Design_Documentation.pdf")
    docx_path = os.path.join(output_dir, "Spectra_Media_Agent_Design_Documentation.docx")
    rtf_path = os.path.join(output_dir, "Spectra_Media_Agent_Design_Documentation.rtf")
    
    print("Generating design documentation...")
    print("-" * 60)
    
    create_pdf(pdf_path)
    create_docx(docx_path)
    create_rtf(rtf_path)
    
    print("-" * 60)
    print("\n✓ All documents generated successfully!")
    print(f"\nOutputs:")
    print(f"  • {pdf_path}")
    print(f"  • {docx_path}")
    print(f"  • {rtf_path}")
